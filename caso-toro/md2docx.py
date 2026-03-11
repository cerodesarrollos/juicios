#!/usr/bin/env python3
"""Convert the denuncia penal MD to a properly formatted DOCX."""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Read the markdown
with open('/tmp/juicios/caso-toro/DENUNCIA-PENAL-FRIGERI-v2.md', 'r') as f:
    content = f.read()

# Remove markdown header
content = content.replace('# FORMULA DENUNCIA. SOLICITA SER PARTE QUERELLANTE.\n', '')

def add_bold_text(paragraph, text):
    """Add text with **bold** markers parsed."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def add_italic_text(paragraph, text):
    """Add text with *italic* markers parsed."""
    parts = re.split(r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)', text)
    for i, part in enumerate(parts):
        if i % 2 == 1:  # italic part
            run = paragraph.add_run(part)
            run.italic = True
        else:
            add_bold_text(paragraph, part)

def process_inline(paragraph, text):
    """Process bold and italic."""
    # First handle bold+italic combos, then bold, then italic
    add_italic_text(paragraph, text)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FORMULA DENUNCIA. SOLICITA SER PARTE QUERELLANTE.')
run.bold = True
run.font.size = Pt(14)

lines = content.split('\n')
i = 0
in_table = False
table_rows = []

while i < len(lines):
    line = lines[i].rstrip()
    
    # Skip horizontal rules
    if line.strip() == '---':
        i += 1
        continue
    
    # Skip empty lines
    if not line.strip():
        i += 1
        continue
    
    # Headers
    if line.startswith('## '):
        header_text = line[3:].strip()
        # Remove markdown formatting from headers
        header_text = header_text.replace('**', '')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(13)
        i += 1
        continue
    
    if line.startswith('### '):
        header_text = line[4:].strip()
        header_text = header_text.replace('**', '')
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(12)
        run.underline = True
        i += 1
        continue
    
    # Tables - collect and create
    if line.strip().startswith('|'):
        table_rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row = lines[i].strip()
            # Skip separator rows
            if re.match(r'\|[\s\-\|:]+\|', row):
                i += 1
                continue
            cells = [c.strip() for c in row.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
        
        if table_rows:
            num_cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=num_cols)
            table.style = 'Table Grid'
            for ri, row_data in enumerate(table_rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < num_cols:
                        cell = table.cell(ri, ci)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        # Clean markdown from cell text
                        clean = cell_text.replace('**', '')
                        p.add_run(clean).font.size = Pt(10)
                        if ri == 0:
                            p.runs[0].bold = True
            doc.add_paragraph()  # spacing after table
        continue
    
    # Numbered lists
    if re.match(r'^\d+\.', line.strip()):
        text = re.sub(r'^\d+\.\s*', '', line.strip())
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        num_match = re.match(r'^(\d+)', line.strip())
        run = p.add_run(f"{num_match.group(1)}. ")
        run.bold = True
        process_inline(p, text)
        i += 1
        continue
    
    # Bullet points
    if line.strip().startswith('- '):
        text = line.strip()[2:]
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.add_run('• ')
        process_inline(p, text)
        # Check for sub-bullets
        i += 1
        while i < len(lines) and lines[i].strip().startswith('- ') and lines[i].startswith('  '):
            sub_text = lines[i].strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(2.5)
            p.add_run('◦ ')
            process_inline(p, sub_text)
            i += 1
        continue
    
    # Regular paragraph - handle the → symbol lines specially
    if line.strip().startswith('→'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        process_inline(p, line.strip())
        i += 1
        continue
    
    # Regular paragraph
    p = doc.add_paragraph()
    text = line.strip()
    
    # Handle indented paragraphs
    if line.startswith('\t') or line.startswith('   '):
        p.paragraph_format.first_line_indent = Cm(1.25)
    
    process_inline(p, text)
    i += 1

# Add signature lines at the end
doc.add_paragraph()
doc.add_paragraph()

sig1 = doc.add_paragraph()
sig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
sig1.add_run('_' * 40)

sig1_label = doc.add_paragraph()
sig1_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
sig1_label.add_run('[Firma del denunciante]')

doc.add_paragraph()

sig2 = doc.add_paragraph()
sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sig2.add_run('_' * 40)

sig2_label = doc.add_paragraph()
sig2_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
sig2_label.add_run('[Firma del letrado patrocinante]')

# Save
output_path = '/tmp/juicios/caso-toro/DENUNCIA-PENAL-FRIGERI-v2.docx'
doc.save(output_path)
print(f"Saved to {output_path}")
