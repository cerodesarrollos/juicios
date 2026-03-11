import json, re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('/tmp/juicios/caso-toro/chunks_all.json') as f:
    data = json.load(f)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2)

ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(11)
ns.paragraph_format.space_after = Pt(3)
ns.paragraph_format.line_spacing = 1.15

# Title page
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('TRANSCRIPCIÓN DE CONVERSACIONES DE WHATSAPP')
r.bold = True; r.font.size = Pt(16)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Y DOCUMENTOS DE EVIDENCIA')
r2.bold = True; r2.font.size = Pt(14)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('CASO: FRIGERI, Matías Juan vs. CHAVES, Franco Maximiliano')
r3.font.size = Pt(12)

t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
t4.add_run('Expediente vinculado: Denuncia Penal por Estafa (Art. 172 CP)')

doc.add_paragraph()
t5 = doc.add_paragraph()
t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
t5.add_run(f'Total de registros: 634 mensajes y documentos')

t6 = doc.add_paragraph()
t6.alignment = WD_ALIGN_PARAGRAPH.CENTER
t6.add_run(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y")}')

doc.add_paragraph()
doc.add_paragraph()

# Legend
leg = doc.add_paragraph()
leg.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = leg.add_run('REFERENCIAS:')
r.bold = True; r.font.size = Pt(10)

for label, desc in [
    ('[EVIDENCIA CLAVE]', 'Mensaje marcado como evidencia relevante para el caso'),
    ('[AUDIO TRANSCRIPTO]', 'Transcripción de mensaje de audio de WhatsApp'),
    ('[documento]', 'Documento o informe adjuntado como evidencia'),
    ('Toro', 'Franco Maximiliano Chaves (demandado) — apodo "Toro"'),
    ('Matias', 'Matías Juan Frigeri Frias (denunciante/actor)'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r1 = p.add_run(f'{label}: ')
    r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

doc.add_page_break()

# Chapter descriptions
ch_names = {
    0: 'DOCUMENTOS E INFORMES DE EVIDENCIA',
    2: 'NEGOCIACIÓN Y COMPRAVENTA DEL AUDI TT (Junio 2025)',
    3: 'DISCUSIONES, FALLAS Y PRESIÓN DE COBRO (Agosto - Octubre 2025)',
    4: 'ACUERDO DE DEVOLUCIÓN DEL VEHÍCULO (Noviembre 2025)',
    5: 'INCUMPLIMIENTO, MENTIRAS Y EVIDENCIA DE FRAUDE (Noviembre 2025 - Febrero 2026)',
}

ch_descriptions = {
    0: 'Informes de BCRA, DNRPA, capturas de Instagram, multas de tránsito y demás documentos que conforman la prueba del caso.',
    2: 'Inicio de la relación comercial, oferta del Audi TT, negociación del precio (5 cuotas de USD 5.000), entrega del vehículo y primeras fallas mecánicas.',
    3: 'Reclamos por cobros anticipados, acumulación de fallas mecánicas, mezcla deliberada de deudas, amenaza velada del 22/10/2025 ("de OTRA MANERA"), presión de cobro con intereses usurarios.',
    4: 'Acorralado por amenazas y fallas, Matias ofrece devolver el auto. Acuerdo: devolución + restitución de USD 17.000 en 3 cuotas. Entrega del vehículo el 13/11/2025. Franco paga solo USD 4.700.',
    5: 'Franco declara que el auto está "destruido" mientras las multas de tránsito y publicaciones de Instagram prueban que circula normalmente. Pretensión extorsiva de USD 15.000. Último contacto 22/02/2026.',
}

# Group by chapter
by_chapter = {}
for c in data:
    ch = c.get('chapter') or 0
    by_chapter.setdefault(ch, []).append(c)

# Table of contents
toc_title = doc.add_paragraph()
r = toc_title.add_run('ÍNDICE')
r.bold = True; r.font.size = Pt(14)
doc.add_paragraph()

for ch in sorted(by_chapter.keys()):
    msgs = by_chapter[ch]
    p = doc.add_paragraph()
    r = p.add_run(f'Capítulo {ch if ch > 0 else "0 (Documentos)"}: ')
    r.bold = True
    p.add_run(f'{ch_names.get(ch, "Sin título")} — {len(msgs)} registros')

doc.add_page_break()

# Chapters
for ch in sorted(by_chapter.keys()):
    msgs = by_chapter[ch]
    
    # Chapter header
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(20)
    r = h.add_run(f'CAPÍTULO {ch if ch > 0 else "0"} — {ch_names.get(ch, "")}')
    r.bold = True; r.font.size = Pt(14)
    
    # Description
    desc = doc.add_paragraph()
    desc.paragraph_format.space_after = Pt(10)
    r = desc.add_run(ch_descriptions.get(ch, ''))
    r.italic = True; r.font.size = Pt(10)
    
    count_line = doc.add_paragraph()
    r = count_line.add_run(f'Total de registros en este capítulo: {len(msgs)}')
    r.bold = True; r.font.size = Pt(10)
    
    doc.add_paragraph()  # spacer
    
    # Messages
    for i, msg in enumerate(msgs):
        content = msg.get('content', '')
        meta = msg.get('metadata') or {}
        sender = meta.get('sender', '')
        is_key = meta.get('is_key', False)
        date_str = meta.get('date', '')
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        
        if ch == 0:
            # Documents - just show content
            if is_key:
                p.paragraph_format.left_indent = Cm(0.5)
                # Yellow-ish highlight effect via bold
                r = p.add_run('⚡ ')
                r.font.size = Pt(10)
            r = p.add_run(content[:500])
            r.font.size = Pt(9)
            if len(content) > 500:
                r2 = p.add_run(f' [...{len(content)} chars total]')
                r2.font.size = Pt(8); r2.italic = True
        else:
            # Chat messages
            # Parse sender and timestamp from content
            # Format: "Sender (timestamp): message"
            
            if is_key:
                p.paragraph_format.left_indent = Cm(0.3)
                # Mark key evidence
                marker = p.add_run('▶ ')
                marker.bold = True
                marker.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            
            # Format timestamp nicely
            if date_str:
                try:
                    dt = datetime.fromisoformat(date_str.replace('+00:00', '+00:00'))
                    nice_date = dt.strftime('%d/%m/%Y %H:%M')
                except:
                    nice_date = ''
            else:
                nice_date = ''
            
            # Extract just the message text (remove "Sender (date): " prefix)
            msg_text = content
            prefix_match = re.match(r'^(Toro|Matias|Sistema)\s*\([^)]+\):\s*', content)
            if prefix_match:
                msg_text = content[prefix_match.end():]
            
            # Sender name
            if sender:
                sr = p.add_run(f'{sender}')
                sr.bold = True
                sr.font.size = Pt(10)
                if sender == 'Toro':
                    sr.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
                elif sender == 'Matias':
                    sr.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                
                if nice_date:
                    dr = p.add_run(f'  [{nice_date}]')
                    dr.font.size = Pt(8)
                    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                
                p.add_run('\n')
            
            # Message content
            mr = p.add_run(msg_text)
            mr.font.size = Pt(10)
            
            # Mark audio transcriptions
            if '[AUDIO' in content or 'audio' in content.lower() and 'transcri' in content.lower():
                mr.italic = True
    
    # Page break between chapters
    if ch < max(by_chapter.keys()):
        doc.add_page_break()

# Footer
doc.add_paragraph()
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run('— FIN DEL DOCUMENTO —')
r.bold = True; r.font.size = Pt(11)

foot2 = doc.add_paragraph()
foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot2.add_run('Este documento contiene la totalidad de los mensajes de WhatsApp\n'
                    'y documentos de evidencia cargados en el sistema Litigium para el caso\n'
                    'FRIGERI vs CHAVES (ID: 571d7e4e-73af-4878-9a8b-05b0a7cd2d49)')
r.font.size = Pt(9); r.italic = True

doc.save('/tmp/juicios/caso-toro/CHAT-COMPLETO-POR-CAPITULO.docx')
import os
sz = os.path.getsize('/tmp/juicios/caso-toro/CHAT-COMPLETO-POR-CAPITULO.docx')
print(f'OK: CHAT-COMPLETO-POR-CAPITULO.docx ({sz:,} bytes)')
